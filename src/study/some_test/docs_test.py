from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import random

def create_numbered_list_style(doc):
    # 创建编号列表样式的XML
    numbering_xml = '''<?xml version="1.0" ?>
    <w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:abstractNum w:abstractNumId="0">
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="720" w:hanging="360"/>
                </w:pPr>
            </w:lvl>
        </w:abstractNum>
        <w:num w:numId="1">
            <w:abstractNumId w:val="0"/>
        </w:num>
        <w:num w:numId="2">
            <w:abstractNumId w:val="0"/>
        </w:num>
    </w:numbering>
    '''
    numbering_part = doc.part.numbering_part
    if not numbering_part:
        numbering_part = doc.part.add_numbering_part()
    element = OxmlElement('w:abstractNum')
    element.set(qn('w:abstractNumId'), '0')
    return element

def adjust_paragraph_indent(paragraph, left_indent=0):
    """调整段落的缩进"""
    p_format = paragraph.paragraph_format
    p_format.left_indent = left_indent  # 设置左缩进
    p_format.first_line_indent = None  # 确保首行没有额外缩进

# 创建文档
doc = Document()

# 设置编号样式
create_numbered_list_style(doc)

# 添加第一个列表 (1-10)
for i in range(10):
    p = doc.add_paragraph()
    p.style = 'List Number'
    p.add_run(f'word{random.randint(1, 100)}')
    adjust_paragraph_indent(p, left_indent=0)  # 确保左对齐

# 添加空行
doc.add_paragraph()

# 添加第二个列表 (1-5)，并调整缩进
for i in range(5):
    p = doc.add_paragraph()
    p.style = 'List Number 2'
    p.add_run(f'word{random.randint(1, 100)}')
    adjust_paragraph_indent(p, left_indent=0)  # 确保左对齐

# 保存文档
doc.save('numbered_lists.docx')
