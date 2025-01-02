from docx import Document
import json

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


def adjust_paragraph_indent(paragraph, left_indent=720):
    """调整段落的缩进"""
    p_format = paragraph.paragraph_format
    p_format.left_indent = left_indent  # 设置左缩进
    p_format.first_line_indent = None  # 确保首行没有额外缩进


def create_numbered_list_style(doc):
    # 创建编号列表样式的XML
    numbering_xml = '''<?xml version="1.0" ?>
    <w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:abstractNum w:abstractNumId="0">
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="720" w:hanging="360"/>
                </w:pPr>
            </w:lvl>
        </w:abstractNum>
        <w:num w:numId="1">
            <w:abstractNumId w:val="0"/>
        </w:num>
        <w:num w:numId="2">
            <w:abstractNumId w:val="0"/>
        </w:num>
    </w:numbering>
    '''
    numbering_part = doc.part.numbering_part
    if not numbering_part:
        numbering_part = doc.part.add_numbering_part()
    element = OxmlElement('w:abstractNum')
    element.set(qn('w:abstractNumId'), '0')
    return element


# 读取JSON文件
def read_json(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def convert_json_2_work(json_data, output_path, book_name):
    # 初始化Word文档
    doc = Document()
    # 设置编号样式
    num_style = create_numbered_list_style(doc)

    # 添加标题
    # title = doc.add_heading('《活着》名著测试题', level=1)
    title = doc.add_paragraph(book_name + '名著测试题')
    # 设置标题段落居中
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 设置字体样式
    run = title.runs[0]  # 获取标题的第一个运行
    run.font.name = '宋体'  # 设置字体为宋体
    run.bold = True  # 加粗

    # 由于中文字体设置需要特殊处理，这里额外设置 eastAsia 字体
    r = run._element
    rPr = r.find(qn('w:rPr'))  # 查找运行的样式属性
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.append(rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)

    rFonts.set(qn('w:eastAsia'), '宋体（正文）')  # 设置东亚字体为宋体

    # 设置字体大小为小三（15磅）
    run.font.size = Pt(15)  # 使用 Pt 对象设置字体大小

    # 添加 "一、选择题" 字段
    paragraph = doc.add_paragraph('一、选择题')

    # 设置段落左对齐
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 设置 "一、选择题" 的字体样式
    run = paragraph.runs[0]
    run.font.name = '宋体'  # 设置字体为宋体
    run.bold = True  # 加粗
    run.font.size = Pt(12)  # 小四对应 12 磅

    # 设置中文字体样式
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.append(rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)

    rFonts.set(qn('w:eastAsia'), '宋体（正文）')  # 设置东亚字体为宋体

    # 遍历 JSON 数据的 knowledge 字段并添加到文档
    for section in json_data['questions']:
        knowledge_field = section['knowledge']  # 获取知识点
        # 添加段落
        paragraph = doc.add_paragraph('【知识点】' + knowledge_field)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.runs[0]
        run.font.name = '宋体（正文）'
        run.bold = True  # 设置加粗
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.append(rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)  # 设置小四字体（12磅）

        # 添加对应的每个问题和选项
        for question in section['questions']:
            # 添加问题为列表项
            question_paragraph = doc.add_paragraph(style='List Number')
            question_paragraph.add_run(f"{question['question']}")
            question_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = question_paragraph.runs[0]
            run.font.name = '宋体'
            run.bold = False
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)  # 小四字体

            # 添加选项
            for option in question['options']:
                option_paragraph = doc.add_paragraph(option)
                option_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                option_run = option_paragraph.runs[0]
                option_run.font.name = '宋体（正文）'
                option_run.bold = False
                r = option_run._element
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.append(rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), '宋体（正文）')
                option_run.font.size = Pt(12)  # 小四字体

    # 添加 "《幸福来临时》名著测试题" 字段
    title_2 = doc.add_paragraph(book_name + '名著测试题')
    title_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    run_2 = title_2.runs[0]
    run_2.font.name = '宋体（正文）'
    run_2.bold = True  # 加粗
    r = run_2._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体（正文）')
    run_2.font.size = Pt(15)  # 小四字体

    # 添加 "参考答案" 字段
    reference_title = doc.add_paragraph('参考答案')
    reference_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    reference_run = reference_title.runs[0]
    reference_run.font.name = '宋体（正文）'
    reference_run.bold = True  # 加粗
    r = reference_run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体（正文）')
    reference_run.font.size = Pt(15)  # 小四字体

    # 添加 "一、选择题" 字段
    paragraph_kk = doc.add_paragraph('一、选择题')

    # 设置段落左对齐
    paragraph_kk.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 设置 "一、选择题" 的字体样式
    run = paragraph_kk.runs[0]
    run.font.name = '宋体'  # 设置字体为宋体
    run.bold = True  # 加粗
    run.font.size = Pt(12)  # 小四对应 12 磅

    # 设置中文字体样式
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.append(rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)

    rFonts.set(qn('w:eastAsia'), '宋体（正文）')  # 设置东亚字体为宋体

    # 创建自定义编号样式
    for section in json_data["questions"]:
        for question in section["questions"]:

            # 添加答案段落
            answer_paragraph = doc.add_paragraph(style='List Number 2')  # 不使用默认样式
            answer_paragraph.add_run(question["answer"])
            answer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            adjust_paragraph_indent(answer_paragraph, left_indent=720)  # 确保左对齐

            run = answer_paragraph.runs[0]
            run.font.name = '宋体'
            run.bold = False
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)  # 小四字体

    doc.save(output_path)


# 主函数
def create_docx(json_file, book_name):
    # json_file = '/Users/macmini/PycharmProjects/LangChainStudy/src/study/三国演义.json'  # JSON文件路径

    output_file = '/Users/macmini/Documents/题目生成/' + book_name +'名著测试题.docx'  # 输出Word文件路径

    # 读取JSON数据
    json_data = read_json(json_file)

    # 创建Word文档
    convert_json_2_work(json_data, output_file, book_name)
    # create_word_from_json(json_data, output_file)
    print(f"Word文档已成功保存到: {output_file}")
    return output_file


if __name__ == "__main__":
    # json_file = '/Users/macmini/PycharmProjects/LangChainStudy/src/study/活着.json'
    # json_data = read_json(json_file)
    create_docx()
